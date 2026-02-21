"""
DOCX Loader — structure-preserving.

Key insight: DOCX files have semantic structure (Heading 1/2/3, body paragraphs,
tables) that PyPDF-style flat extraction throws away. We preserve it.

What we extract:
- Headings with their level (h1, h2, h3) → attached to subsequent paragraphs
  as metadata so the structural chunker can use section boundaries
- Tables → serialized row-by-row with column headers, same as PDF tables
- Body paragraphs → grouped under their nearest heading

Why this matters for RAG:
A policy document with "Section 3: Parental Leave > 3.2 Paternity Leave > ..."
becomes retrievable by section, not just keyword. The heading metadata lets the
model say "this chunk is from the Paternity Leave section" in its answer.
"""
from pathlib import Path

from langchain_core.documents import Document

from app.loaders.base import BaseLoader


class DocxLoader(BaseLoader):
    SUPPORTED_EXTENSIONS = [".docx", ".doc"]

    def load(self, file_path: Path) -> list[Document]:
        import docx
        doc = docx.Document(str(file_path))
        documents = []
        current_section = ""
        current_heading_level = 0
        buffer = []

        def flush_buffer():
            if buffer:
                documents.append(Document(
                    page_content="\n".join(buffer),
                    metadata={
                        "source": file_path.name,
                        "section": current_section,
                        "heading_level": current_heading_level,
                    },
                ))
                buffer.clear()

        for para in doc.paragraphs:
            style = (para.style.name or "").strip()
            text = para.text.strip()
            if not text:
                continue

            if style.startswith("Heading"):
                flush_buffer()
                try:
                    level = int(style.split()[-1])
                except (ValueError, IndexError):
                    level = 1
                current_section = text
                current_heading_level = level
                documents.append(Document(
                    page_content=text,
                    metadata={
                        "source": file_path.name,
                        "section": text,
                        "heading_level": level,
                        "is_heading": True,
                    },
                ))
            else:
                buffer.append(text)

        for table in doc.tables:
            flush_buffer()
            serialized = self._serialize_docx_table(table)
            if serialized:
                documents.append(Document(
                    page_content=serialized,
                    metadata={
                        "source": file_path.name,
                        "section": current_section,
                        "is_table": True,
                    },
                ))

        flush_buffer()
        return documents

    def _serialize_docx_table(self, table) -> str:
        rows = []
        for row in table.rows:
            rows.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "TABLE:\n" + "\n".join(rows) if rows else ""
